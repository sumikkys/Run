from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document


DEFAULT_DOCX = Path(r"C:\Users\Yifan\Desktop\北京换乘人工标注数据（需清洗）.docx")
REPO_ROOT = Path(__file__).resolve().parents[2]
DEFAULT_OUTPUT = REPO_ROOT / "data" / "manual_transfers.json"


def compact(text: str) -> str:
    return re.sub(r"\s+", " ", str(text or "")).strip()


def strip_index(place: str) -> str:
    return re.sub(r"^\d+\s*", "", compact(place))


def estimate_minutes(text: str, target: str) -> tuple[float, float, str]:
    merged = f"{target} {text}"
    explicit = explicit_minutes(merged)
    if explicit:
        p50 = min(explicit)
        p90 = max(explicit) if max(explicit) > p50 else p50 + 2
        return float(p50), float(max(p90, p50 + 1)), "explicit_or_near_explicit"

    if any(word in merged for word in ["同台", "对面站台", "直接走到对面"]):
        return 2.0, 4.0, "heuristic_same_platform"
    if any(word in merged for word in ["直达", "非常方便", "很快", "直接到达", "下楼梯就是"]):
        return 3.0, 6.0, "heuristic_short_transfer"
    if any(word in merged for word in ["长长", "较长", "绕", "800米", "不互通"]):
        return 8.0, 15.0, "heuristic_long_transfer"
    if any(word in merged for word in ["进站", "安检", "候车大厅", "检票口"]):
        return 8.0, 15.0, "heuristic_station_entry"
    return 5.0, 10.0, "heuristic_default"


def explicit_minutes(text: str) -> list[float]:
    values: list[float] = []
    for match in re.finditer(r"(\d+(?:\.\d+)?)\s*[-到至]\s*(\d+(?:\.\d+)?)\s*分钟", text):
        values.extend([float(match.group(1)), float(match.group(2))])
    for match in re.finditer(r"(\d+(?:\.\d+)?)\s*(?:分|分钟)", text):
        values.append(float(match.group(1)))
    if "一两分钟" in text:
        values.extend([1.0, 2.0])
    if "不到一分钟" in text:
        values.append(1.0)
    if "半小时" in text:
        values.append(30.0)
    return values


def extract_docx(docx_path: Path) -> dict[str, object]:
    doc = Document(docx_path)
    entries: list[dict[str, object]] = []

    for table_index, table in enumerate(doc.tables):
        for row_index, row in enumerate(table.rows[1:], start=1):
            cells = [compact(cell.text) for cell in row.cells]
            if len(cells) < 3 or not cells[0] or not cells[1]:
                continue
            place = strip_index(cells[0])
            target = cells[1]
            advice = cells[2]
            p50, p90, estimate_source = estimate_minutes(advice, target)
            entries.append(
                {
                    "place": place,
                    "target": target,
                    "p50_min": p50,
                    "p90_min": p90,
                    "advice": advice,
                    "source": "manual_label_docx",
                    "estimate_source": estimate_source,
                    "confidence": "medium" if estimate_source.startswith("explicit") else "low",
                    "doc_table": table_index,
                    "doc_row": row_index,
                }
            )

    entries.extend(extract_paragraph_route_entries(doc.paragraphs))
    return {
        "source_docx": str(docx_path),
        "entry_count": len(entries),
        "entries": entries,
    }


def extract_paragraph_route_entries(paragraphs) -> list[dict[str, object]]:
    entries: list[dict[str, object]] = []
    current_place = ""
    for paragraph in paragraphs:
        text = compact(paragraph.text)
        if not text:
            continue
        header = re.match(r"^\|\s*\d+\s*\|\s*\*\*(.+?)\*\*", text)
        if header:
            current_place = header.group(1)
            continue
        if not current_place:
            continue
        if "：" not in text and ":" not in text:
            continue
        target, advice = re.split(r"：|:", text, maxsplit=1)
        if not any(token in target for token in ["->", "-", "⇌", "换乘", "进站"]):
            continue
        p50, p90, estimate_source = estimate_minutes(advice, target)
        entries.append(
            {
                "place": current_place,
                "target": target,
                "p50_min": p50,
                "p90_min": p90,
                "advice": advice,
                "source": "manual_label_docx_paragraph",
                "estimate_source": estimate_source,
                "confidence": "medium" if estimate_source.startswith("explicit") else "low",
            }
        )
    return entries


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", default=str(DEFAULT_DOCX))
    parser.add_argument("--output", default=str(DEFAULT_OUTPUT))
    args = parser.parse_args()

    payload = extract_docx(Path(args.docx))
    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with output_path.open("w", encoding="utf-8") as file:
        json.dump(payload, file, ensure_ascii=False, indent=2)
    print(json.dumps({"output": str(output_path), "entry_count": payload["entry_count"]}, ensure_ascii=False))


if __name__ == "__main__":
    main()

